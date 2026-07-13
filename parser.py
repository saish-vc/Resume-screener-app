import io
from pypdf import PdfReader
import docx

def extract_text_from_pdf(file_stream: io.BytesIO) -> str:
    """
    Extracts text from a PDF file stream using pypdf.
    """
    try:
        reader = PdfReader(file_stream)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")

def extract_text_from_docx(file_stream: io.BytesIO) -> str:
    """
    Extracts text from a DOCX file stream using python-docx.
    Also extracts text from tables inside the DOCX.
    """
    try:
        doc = docx.Document(file_stream)
        text_parts = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
                        
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error parsing DOCX file: {str(e)}")

def extract_text_from_txt(file_stream: io.BytesIO) -> str:
    """
    Extracts text from a plain text file stream, trying UTF-8 first, then falling back to Latin-1.
    """
    try:
        # Seek to start to be safe
        file_stream.seek(0)
        content = file_stream.read()
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")
    except Exception as e:
        raise ValueError(f"Error parsing TXT file: {str(e)}")

def parse_resume(file_stream: io.BytesIO, filename: str) -> str:
    """
    Determines the file type based on its extension and extracts text.
    Raises ValueError for unsupported formats or parsing errors.
    """
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    
    # Ensure stream is at the beginning
    file_stream.seek(0)
    
    if ext == "pdf":
        return extract_text_from_pdf(file_stream)
    elif ext in ["docx", "doc"]:
        # python-docx requires docx, .doc is legacy binary but sometimes renamed or partially compatible.
        # We will attempt to parse .docx, and notify on exception.
        return extract_text_from_docx(file_stream)
    elif ext == "txt":
        return extract_text_from_txt(file_stream)
    else:
        raise ValueError(f"Unsupported file format '.{ext}'. Supported formats: .pdf, .docx, .txt")
