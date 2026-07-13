import os
import docx

def create_txt_resume(filepath):
    content = """JANE SMITH
Email: jane.smith@example.com
Phone: +1-555-456-7890
Location: San Francisco, CA
GitHub: github.com/janesmith

PROFESSIONAL SUMMARY
Dedicated Python Backend Developer with 3 years of experience. Specialized in designing, building, and optimizing server-side APIs and database systems. Experienced in working in collaborative team environments with a strong focus on clean code and performance.

CORE SKILLS
- Languages: Python, SQL, Bash
- Frameworks: Flask, FastAPI
- Databases: PostgreSQL, SQLite, Redis
- Tools: Git, Linux, Nginx, Docker
- Soft Skills: Problem Solving, Teamwork, Analytical

PROFESSIONAL EXPERIENCE
Software Engineer | Backend Systems
DataGenius Corp (2024 - Present)
- Developed and maintained critical backend APIs using Python and Flask, serving over 100k daily users.
- Designed database schemas and optimized SQL queries in PostgreSQL, reducing query latency by 40%.
- Integrated third-party APIs and built custom middleware to handle real-time data sync.
- Utilized Git for version control and participated in daily standups and agile sprints.

Junior Python Developer
WebStart Inc (2023 - 2024)
- Assisted in building server-side logic and endpoints using FastAPI.
- Wrote unit tests and integration tests to ensure 90%+ code coverage.
- Debugged backend application errors and collaborated with frontend developers.

EDUCATION
B.S. in Computer Science
University of California, Berkeley (Graduated 2023)
"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content.strip())
    print(f"Created plain text resume: {filepath}")

def create_john_doe_docx(filepath):
    doc = docx.Document()
    
    # Title/Header
    title = doc.add_paragraph()
    run = title.add_run("JOHN DOE")
    run.font.size = docx.shared.Pt(24)
    run.bold = True
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.add_run("Email: john.doe@email.com | Phone: (555) 019-2834 | New York, NY\nLinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe")
    contact.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    p = doc.add_paragraph(
        "Highly skilled and results-oriented Senior Full-Stack Developer with over 6 years of experience "
        "designing, building, and deploying scalable web applications. Expert in Python and React, with a "
        "deep focus on clean code, responsive UI, database design, and cloud containerization."
    )
    
    # Skills
    doc.add_heading("Core Technical Skills", level=1)
    p_skills = doc.add_paragraph()
    p_skills.add_run("Languages: ").bold = True
    p_skills.add_run("Python, JavaScript, TypeScript, SQL, HTML, CSS\n")
    p_skills.add_run("Backend Frameworks: ").bold = True
    p_skills.add_run("Django, FastAPI, Flask, Node.js\n")
    p_skills.add_run("Frontend Frameworks: ").bold = True
    p_skills.add_run("React, Redux, Tailwind, Bootstrap\n")
    p_skills.add_run("Databases & Tools: ").bold = True
    p_skills.add_run("PostgreSQL, MongoDB, MySQL, Git, Docker, Kubernetes\n")
    p_skills.add_run("Cloud & DevOps: ").bold = True
    p_skills.add_run("AWS (Amazon Web Services), Jenkins, CI/CD, Linux\n")
    p_skills.add_run("Methodologies: ").bold = True
    p_skills.add_run("Agile, Scrum, Unit Testing, System Architecture")
    
    # Experience
    doc.add_heading("Professional Experience", level=1)
    
    # Role 1
    doc.add_heading("Senior Full-Stack Developer | CloudTech Solutions (2023 - Present)", level=2)
    p_exp1 = doc.add_paragraph(
        "- Led a team of 4 developers to build a modern SaaS analytics platform using Python, Django, and React.\n"
        "- Containerized applications using Docker and deployed them to AWS ECS, automating deployments via Jenkins CI/CD pipelines.\n"
        "- Optimized SQL database schemas and PostgreSQL query performance, resulting in a 50% speedup in reports loading.\n"
        "- Championed Agile software development practices, hosting sprint planning and review meetings.\n"
        "- Exercised excellent communication skills while interacting with business analysts and product managers to translate requirements into technical designs."
    )
    
    # Role 2
    doc.add_heading("Full-Stack Software Engineer | DevBuilders Inc (2020 - 2023)", level=2)
    p_exp2 = doc.add_paragraph(
        "- Designed and implemented robust REST APIs using FastAPI and Flask in Python.\n"
        "- Created modern and highly interactive frontend dashboards using React and TypeScript.\n"
        "- Managed source code using Git and GitHub, performing rigorous code reviews to maintain code quality.\n"
        "- Implemented unit testing and integration testing frameworks, increasing coverage to 85%."
    )
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.S. in Software Engineering\nRochester Institute of Technology (Graduated 2020)")
    
    doc.save(filepath)
    print(f"Created docx resume (John Doe): {filepath}")

def create_bob_johnson_docx(filepath):
    doc = docx.Document()
    
    # Title/Header
    title = doc.add_paragraph()
    run = title.add_run("BOB JOHNSON")
    run.font.size = docx.shared.Pt(24)
    run.bold = True
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.add_run("Email: bob.johnson@workmail.net | Phone: 555.223.4456 | Chicago, IL")
    contact.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    p = doc.add_paragraph(
        "Dynamic, results-driven Sales Manager with 8+ years of experience leading high-performing sales teams, "
        "retaining key client accounts, and exceeding corporate revenue targets. Expert in strategic negotiation, "
        "relationship management, and presentation."
    )
    
    # Skills
    doc.add_heading("Core Professional Skills", level=1)
    p_skills = doc.add_paragraph()
    p_skills.add_run("Core Competencies: ").bold = True
    p_skills.add_run("Sales, Sales Management, Account Management, Negotiation, Presentation, CRM (Salesforce)\n")
    p_skills.add_run("Leadership Skills: ").bold = True
    p_skills.add_run("Team Leadership, Communication, Project Management, Collaboration, Creativity, Scrum")
    
    # Experience
    doc.add_heading("Professional Experience", level=1)
    
    # Role 1
    doc.add_heading("Senior Sales Manager | Velocity Retail Group (2022 - Present)", level=2)
    p_exp1 = doc.add_paragraph(
        "- Directed a sales team of 12 reps, exceeding the annual sales quota by 18% in consecutive fiscal years.\n"
        "- Managed top-tier enterprise accounts, maintaining a customer retention rate of 95%.\n"
        "- Conducted training sessions on advanced sales negotiation and project management tools.\n"
        "- Collaborated across departments to design customized pricing strategies and customer offers."
    )
    
    # Role 2
    doc.add_heading("Account Executive | Summit Enterprises (2018 - 2022)", level=2)
    p_exp2 = doc.add_paragraph(
        "- Acquired 45 new business contracts, expanding market presence in the Midwest region.\n"
        "- Delivered high-impact sales presentations to C-suite executives, closing key deals.\n"
        "- Actively resolved customer complaints and refined client onboarding processes."
    )
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.A. in Business Administration\nDePaul University (Graduated 2018)")
    
    doc.save(filepath)
    print(f"Created docx resume (Bob Johnson): {filepath}")

if __name__ == "__main__":
    os.makedirs("sample_resumes", exist_ok=True)
    create_txt_resume("sample_resumes/jane_smith_resume.txt")
    create_john_doe_docx("sample_resumes/john_doe_resume.docx")
    create_bob_johnson_docx("sample_resumes/bob_johnson_resume.docx")
    print("All sample resumes successfully generated!")
